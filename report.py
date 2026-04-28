
from __future__ import annotations

import io
from pathlib import Path
from typing import Any

import cv2
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt

def _image_to_stream(image: np.ndarray) -> io.BytesIO:
    ok, buffer = cv2.imencode(".png", image)
    if not ok:
        raise RuntimeError("Failed to encode image for report embedding.")
    return io.BytesIO(buffer.tobytes())

def _histogram_stream(image: np.ndarray, title: str) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(4.2, 3.0), dpi=160)

    if image.ndim == 2:
        ax.hist(image.reshape(-1), bins=256, range=(0, 256), histtype="step", linewidth=1.4, color="black")
        ax.set_xlabel("Intensity")
        ax.set_ylabel("Count")
    else:
        channels = cv2.split(image)
        labels = ["B", "G", "R"]
        for channel, label in zip(channels, labels):
            ax.hist(channel.reshape(-1), bins=256, range=(0, 256), histtype="step", linewidth=1.1, label=label)
        ax.set_xlabel("Intensity")
        ax.set_ylabel("Count")
        ax.legend(frameon=False, fontsize=8)

    ax.set_title(title, fontsize=10)
    ax.set_xlim(0, 255)
    ax.tick_params(labelsize=8)
    fig.tight_layout()

    stream = io.BytesIO()
    fig.savefig(stream, format="png", bbox_inches="tight")
    plt.close(fig)
    stream.seek(0)
    return stream

def _set_cell_text(cell, lines: list[str], bold_first: bool = True) -> None:
    cell.text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9)
        if idx == 0 and bold_first:
            run.bold = True

def generate_report(records: list[dict[str, Any]], report_path: str | Path, title: str) -> Path:
    report_path = Path(report_path)
    report_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

    doc.add_heading(title, level=0)

    for rec in records:
        doc.add_heading(rec["name"], level=1)

        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"

        _set_cell_text(table.cell(0, 0), [rec["left_title"]])
        _set_cell_text(table.cell(0, 1), [rec["right_title"]])

        p0 = table.cell(1, 0).paragraphs[0]
        p0.alignment = 1
        p0.add_run().add_picture(_image_to_stream(rec["left_image"]), width=Inches(2.7))

        p1 = table.cell(1, 1).paragraphs[0]
        p1.alignment = 1
        p1.add_run().add_picture(_image_to_stream(rec["right_image"]), width=Inches(2.7))

        h0 = table.cell(2, 0).paragraphs[0]
        h0.alignment = 1
        h0.add_run().add_picture(_histogram_stream(rec["left_image"], f"{rec['left_title']} Histogram"), width=Inches(2.7))

        h1 = table.cell(2, 1).paragraphs[0]
        h1.alignment = 1
        h1.add_run().add_picture(_histogram_stream(rec["right_image"], f"{rec['right_title']} Histogram"), width=Inches(2.7))

        doc.add_paragraph("")

        metric_table = doc.add_table(rows=2, cols=2)
        metric_table.style = "Table Grid"

        _set_cell_text(
            metric_table.cell(0, 0),
            [
                f"Entropy: {rec['left_entropy']:.4f}",
                f"Corr H: {rec['left_corr'][0]:.4f}",
                f"Corr V: {rec['left_corr'][1]:.4f}",
                f"Corr D: {rec['left_corr'][2]:.4f}",
            ],
        )
        _set_cell_text(
            metric_table.cell(0, 1),
            [
                f"Entropy: {rec['right_entropy']:.4f}",
                f"Corr H: {rec['right_corr'][0]:.4f}",
                f"Corr V: {rec['right_corr'][1]:.4f}",
                f"Corr D: {rec['right_corr'][2]:.4f}",
            ],
        )

        merged = metric_table.cell(1, 0).merge(metric_table.cell(1, 1))
        _set_cell_text(merged, [rec["extra_line1"], rec["extra_line2"], rec["extra_line3"]])

        doc.add_paragraph("")

    doc.save(str(report_path))
    return report_path
